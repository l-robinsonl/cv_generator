"""Render model output into real TXT, DOCX, and PDF files in memory."""

from io import BytesIO
import html
import re

from docx import Document
from docx.shared import Inches, Pt
from reportlab.lib.enums import TA_CENTER
from reportlab.lib.pagesizes import A4, LETTER
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer


MIME_TYPES = {
    "txt": "text/plain",
    "docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    "pdf": "application/pdf",
}


def _lines(content: str):
    for raw_line in content.replace("\r\n", "\n").split("\n"):
        line = raw_line.strip()
        if not line:
            yield "blank", ""
        elif line.startswith("### "):
            yield "heading3", line[4:].strip()
        elif line.startswith("## "):
            yield "heading2", line[3:].strip()
        elif line.startswith("# "):
            yield "heading1", line[2:].strip()
        elif re.match(r"^[-*•]\s+", line):
            yield "bullet", re.sub(r"^[-*•]\s+", "", line)
        else:
            yield "body", line


def _plain(text: str) -> str:
    return re.sub(r"(\*\*|__|`)", "", text)


def render_txt(content: str) -> bytes:
    return content.encode("utf-8")


def render_docx(content: str) -> bytes:
    document = Document()
    section = document.sections[0]
    section.top_margin = Inches(0.65)
    section.bottom_margin = Inches(0.65)
    section.left_margin = Inches(0.7)
    section.right_margin = Inches(0.7)
    document.styles["Normal"].font.name = "Arial"
    document.styles["Normal"].font.size = Pt(10.5)

    for kind, text in _lines(content):
        text = _plain(text)
        if kind == "blank":
            continue
        if kind.startswith("heading"):
            level = int(kind[-1])
            document.add_heading(text, level=min(level, 3))
        elif kind == "bullet":
            document.add_paragraph(text, style="List Bullet")
        else:
            document.add_paragraph(text)

    output = BytesIO()
    document.save(output)
    return output.getvalue()


def render_pdf(content: str, country: str) -> bytes:
    output = BytesIO()
    page_size = LETTER if country == "US" else A4
    document = SimpleDocTemplate(
        output,
        pagesize=page_size,
        rightMargin=0.65 * inch,
        leftMargin=0.65 * inch,
        topMargin=0.6 * inch,
        bottomMargin=0.6 * inch,
        title="Synthetic CV",
    )
    styles = getSampleStyleSheet()
    styles.add(
        ParagraphStyle(
            "CVTitle",
            parent=styles["Title"],
            fontName="Helvetica-Bold",
            fontSize=18,
            leading=21,
            alignment=TA_CENTER,
            spaceAfter=10,
        )
    )
    styles.add(
        ParagraphStyle(
            "CVHeading",
            parent=styles["Heading2"],
            fontName="Helvetica-Bold",
            fontSize=12,
            leading=14,
            spaceBefore=7,
            spaceAfter=4,
        )
    )
    body = ParagraphStyle(
        "CVBody",
        parent=styles["BodyText"],
        fontName="Helvetica",
        fontSize=9.5,
        leading=12,
        spaceAfter=4,
    )
    bullet = ParagraphStyle(
        "CVBullet",
        parent=body,
        leftIndent=12,
        firstLineIndent=-8,
        bulletIndent=2,
    )

    story = []
    for kind, text in _lines(content):
        safe = html.escape(_plain(text)).replace("—", "-").replace("–", "-")
        if kind == "blank":
            story.append(Spacer(1, 3))
        elif kind == "heading1":
            story.append(Paragraph(safe, styles["CVTitle"]))
        elif kind in {"heading2", "heading3"}:
            story.append(Paragraph(safe, styles["CVHeading"]))
        elif kind == "bullet":
            story.append(Paragraph(f"&bull;&nbsp; {safe}", bullet))
        else:
            story.append(Paragraph(safe, body))

    document.build(story)
    return output.getvalue()


def render_document(content: str, output_format: str, country: str) -> bytes:
    if output_format == "txt":
        return render_txt(content)
    if output_format == "docx":
        return render_docx(content)
    if output_format == "pdf":
        return render_pdf(content, country)
    raise ValueError(f"Unsupported output format: {output_format}")
