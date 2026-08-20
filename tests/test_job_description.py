from io import BytesIO
import unittest

from docx import Document

from cv_generator.documents import render_document
from cv_generator.job_description import extract_job_description


class JobDescriptionTests(unittest.TestCase):
    def test_extracts_utf8_text(self):
        text = extract_job_description("role.txt", "Senior seller\nLondon".encode())
        self.assertEqual(text, "Senior seller\nLondon")

    def test_extracts_docx_paragraphs(self):
        document = Document()
        document.add_heading("Marketing Manager", 1)
        document.add_paragraph("Own campaign strategy")
        output = BytesIO()
        document.save(output)

        text = extract_job_description("role.docx", output.getvalue())

        self.assertIn("Marketing Manager", text)
        self.assertIn("Own campaign strategy", text)

    def test_extracts_text_based_pdf(self):
        pdf = render_document(
            "# Sales Director\n\n## Requirements\n\n- Enterprise sales", "pdf", "UK"
        )

        text = extract_job_description("role.pdf", pdf)

        self.assertIn("Sales Director", text)
        self.assertIn("Enterprise sales", text)

    def test_rejects_empty_files(self):
        with self.assertRaisesRegex(ValueError, "does not contain readable text"):
            extract_job_description("role.txt", b"")

    def test_rejects_unsupported_files(self):
        with self.assertRaisesRegex(ValueError, "Unsupported"):
            extract_job_description("role.csv", b"title,description")


if __name__ == "__main__":
    unittest.main()
